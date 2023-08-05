"""
PillowScreenX is a powerful screenshot capturing tool
that leverages the power of the Pillow library to provide high-quality screenshot captures.
With PillowScreenX, you can take full screenshots or capture specific regions
of your screen with ease. The module also provides options to store your screenshots
in a zip file or in a separate folder. Whether you need to capture screenshots for documentation,
debugging, or any other purpose, PillowScreenX is the ideal solution.
"""
import inspect
import os
import re
import time
try:
    import docx
except ImportError:
    print("docx is not installed. Please install it by running the command: pip install python-docx")
try:
    from PIL import ImageGrab
except ImportError:
    print("Pillow is not installed. Please install it by running the command: pip install Pillow")


class PillowScreenX:
    """
    PillowScreenX is a powerful screenshot capturing tool
    that leverages the power of the Pillow library to provide high-quality screenshot captures.

    The default output folder name is the current date and time in the format YYYYMMDD_HHMMSS.
    The default screenshot name is the name of the python file in which this module is called.
    The default screenshot quality is 95.

    Note:
        The output folder name will be read from the environment variable OUTPUT_FOLDER_NAME.
        So, you can set the output folder name by setting the environment variable OUTPUT_FOLDER_NAME.

    Attributes:
        OUTPUT_FOLDER_NAME is OS environment variable OUTPUT_FOLDER_NAME or
                current date and time in the format YYYYMMDD_HHMMSS
    """
    if os.environ.get('OUTPUT_FOLDER_NAME') is not None:
        OUTPUT_FOLDER_NAME = os.environ.get('OUTPUT_FOLDER_NAME')
    else:
        OUTPUT_FOLDER_NAME = time.strftime("%Y%m%d_%H%M%S")

    @classmethod
    def __generate_word_docx(cls, **kwargs) -> None:
        """
        Generate a word document with the screenshots
        """
        WORD_DOCX_NAME = os.environ.get('WORD_DOCX_NAME')
        WORD_DOCX_PATH = os.environ.get('WORD_DOCX_PATH')

        HEADING = kwargs.get('HEADING')
        SUBHEADING = kwargs.get('SUBHEADING')
        SUBHEADING_2 = kwargs.get('SUBHEADING_2')
        IMAGE = kwargs.get('IMAGE')

        # Check if the Word document already exists in the given path
        if os.path.exists(os.path.join(WORD_DOCX_PATH, WORD_DOCX_NAME)):
            # If the document already exists, open it and add the data
            doc = docx.Document(os.path.join(WORD_DOCX_PATH, WORD_DOCX_NAME))
            SUBHEADING_ = SUBHEADING + " | " + SUBHEADING_2
            doc.add_heading(SUBHEADING_, level=2)
            doc.add_picture(IMAGE, width=docx.shared.Inches(6.5))
        else:
            # If the document does not exist, create a new one and add the data
            doc = docx.Document()
            SUBHEADING_ = SUBHEADING + " | " + SUBHEADING_2
            doc.add_heading(SUBHEADING_, level=2)
            doc.add_picture(IMAGE, width=docx.shared.Inches(6.5))

        # Save the document to the given path
        doc.save(os.path.join(WORD_DOCX_PATH, WORD_DOCX_NAME))

    @classmethod
    def __remove_special_chars(cls, string: str) -> str:
        """
        Remove special characters from the given string

        Args:
            string (str): The string to remove special characters

        Returns:
            str: The string without special characters
        """
        clean_string = re.sub(r'\W+', ' ', re.sub(r'_', ' ', string))
        return clean_string

    @classmethod
    def screenshot(cls, name: str = None, wait_time: float = 0.0, **kwargs) -> str:
        """
        Take the screenshot of the screen.

        Args:
            name (str, optional): The name of the screenshot. Defaults to None.
            wait_time (float, optional): The time to wait before taking the screenshot.
                Defaults to 0.0.
            **kwargs: The keyword arguments.

        Keyword Args:
            CAPTURE_SCREENSHOT = False
            FETCH_SCREENSHOT_PATH = True
        """
        SCREENSHOT_QUALITY = 95

        if kwargs.get('CAPTURE_SCREENSHOT') is False:
            CAPTURE_SCREENSHOT = False
        else:
            CAPTURE_SCREENSHOT = True

        if kwargs.get('FETCH_SCREENSHOT_PATH') is False:
            FETCH_SCREENSHOT_PATH = False
        else:
            FETCH_SCREENSHOT_PATH = True

        # Wait time before taking screenshot
        time.sleep(wait_time)

        # Take ScreenShot and store it in bytes format
        screenshot = ImageGrab.grab()

        calling_frame = inspect.stack()[1]
        current_file = calling_frame.filename

        # Final component of a pathname
        file_name = os.path.basename(current_file)

        # Using regex to get the file name without extension
        # The extension length is UNKNOWN
        file_name = re.sub(r'\.\w+$', '', file_name)

        # File directory
        file_location = os.path.dirname(os.path.abspath(current_file))

        # Parent directory of the file
        file_parent_dir = os.path.basename(file_location)

        # Grandparent directory of the file
        file_grandpar_dir = os.path.basename(os.path.dirname(file_location))

        # Get only folders in {file_location} directory
        old_folders_in_file_dir = [f for f in os.listdir(file_location) \
                                        if os.path.isdir(os.path.join(file_location, f))]

        # Delete the folders and its contents in {old_folders_in_file_dir} list
        for folder in old_folders_in_file_dir:
            # Check folder name is in YYYYMMDD_HHMMSS format using regex
            if folder != cls.OUTPUT_FOLDER_NAME and re.match(r'\d{8}_\d{6}', folder):
                for file in os.listdir(f'{file_location}/{folder}'):
                    os.remove(f'{file_location}/{folder}/{file}')
                os.rmdir(f'{file_location}/{folder}')

        output_directory = f'{file_location}\{cls.OUTPUT_FOLDER_NAME}'

        # Create a folder named USE_CASE_FOLDER in the same directory as this file
        # If it doesn't already exist
        if not os.path.exists(f'{output_directory}'):
            os.makedirs(f'{output_directory}')

        if name is None:
            test_case_name = f'{file_name}'
        else:
            test_case_name = name

        # Save the screenshot to the USE_CASE_FOLDER folder
        output_path = f'{output_directory}\{test_case_name}.png'

        # if the file already exists, then add incrementing number to the file name
        if os.path.exists(output_path):
            # get the count of same file name
            file_count = len([f for f in os.listdir(f'{output_directory}') \
                        if f.startswith(file_name)]) + 1

            output_path = f'{output_directory}\{test_case_name}_({file_count}).png'

        if CAPTURE_SCREENSHOT:
            screenshot.save(output_path, quality=SCREENSHOT_QUALITY)
            print(f'\nScreenshot is saved at: {output_path}\n')

            word_meta_data = {
                "HEADING": cls.__remove_special_chars(file_grandpar_dir),
                "SUBHEADING": cls.__remove_special_chars(file_parent_dir),
                "SUBHEADING_2": cls.__remove_special_chars(test_case_name),
                "IMAGE": output_path
            }
            cls.__generate_word_docx(**word_meta_data)

        if FETCH_SCREENSHOT_PATH:
            return output_path
